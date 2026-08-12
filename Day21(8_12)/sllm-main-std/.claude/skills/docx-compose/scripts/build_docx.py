#!/usr/bin/env python3
"""build_docx.py — 마크다운 원고 + PNG 다이어그램을 .docx 로 조립한다.

    python build_docx.py <manuscript.md> <out.docx> [--images-dir DIR]

마크다운 전체를 지원하지 않는다. 템플릿2 보고서가 실제로 쓰는 요소만 처리한다:

    # ## ###          제목 1~3단계
    본문 문단
    - / • 불릿        (들여쓰기 2칸 = 2단계)
    | 표 |            헤더 구분줄이 있는 파이프 표
    ```code```        고정폭 블록
    ![](img.png)      이미지. 다음 줄이 `> 캡션` 이면 캡션으로 붙음
    ---               구분선 (쪽 나눔)
    **굵게** `코드`   문단 안 인라인 서식

한글 폰트는 맑은 고딕으로 고정한다. python-docx 의 rPr 은 East Asian 글꼴을
따로 지정해야 한글이 의도한 서체로 나온다 — w:eastAsia 속성을 손으로 넣는 이유다.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BODY_FONT = "맑은 고딕"
CODE_FONT = "D2Coding"          # 없으면 Word 가 Consolas 로 떨어진다
BODY_SIZE = Pt(10.5)
MAX_IMG_CM = 16.0               # A4 기본 여백 기준 본문 폭


def set_font(run, name=BODY_FONT, size=None, bold=None, color=None):
    run.font.name = name
    if size:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(p, text_, size=BODY_SIZE, base_bold=False):
    """**굵게** 와 `코드` 만 처리한다. 중첩은 지원하지 않는다 —
    보고서 원고에서 중첩 서식이 필요했던 적이 없고, 파서만 복잡해진다."""
    for part in INLINE.split(text_):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            set_font(p.add_run(part[2:-2]), size=size, bold=True)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = p.add_run(part[1:-1])
            set_font(r, name=CODE_FONT, size=Pt(size.pt - 0.5))
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            set_font(p.add_run(part), size=size, bold=base_bold)


def style_heading(doc, text_, level):
    h = doc.add_heading(level=level)
    sizes = {1: Pt(18), 2: Pt(14), 3: Pt(11.5)}
    colors = {1: RGBColor(0x1A, 0x1A, 0x1A), 2: RGBColor(0x1F, 0x4E, 0x79),
              3: RGBColor(0x2E, 0x2E, 0x2E)}
    r = h.add_run(text_)
    set_font(r, size=sizes.get(level, Pt(11)), bold=True, color=colors.get(level))
    h.paragraph_format.space_before = Pt(16 if level <= 2 else 10)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_body(doc, text_):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.5
    add_inline(p, text_)
    return p


def add_bullet(doc, text_, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    add_inline(p, text_)
    return p


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("\n".join(lines))
    set_font(r, name=CODE_FONT, size=Pt(9))
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # 회색 배경 — 코드가 본문에 섞여 읽히는 걸 막는다
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F4F4F4")
    p._p.get_or_add_pPr().append(shd)
    return p


def add_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.autofit = True
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j >= len(t.rows[i].cells):
                continue
            c = t.rows[i].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, cell, size=Pt(9.5), base_bold=(i == 0))
            if i == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "EDF2F7")
                c._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def add_image(doc, path, caption=None, max_cm=MAX_IMG_CM):
    if not path.exists():
        print(f"   !! 이미지 없음: {path}", file=sys.stderr)
        return
    from PIL import Image  # noqa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(max_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(12)
        r = cp.add_run(caption)
        set_font(r, size=Pt(9), color=RGBColor(0x77, 0x77, 0x77))


def image_size_cm(path, max_cm=MAX_IMG_CM, max_h_cm=13.0):
    """세로로 긴 그림이 한 쪽을 통째로 잡아먹지 않게 높이로도 제한한다."""
    try:
        from PIL import Image
        with Image.open(path) as im:
            w, h = im.size
        by_h = max_h_cm * w / h
        return min(max_cm, by_h)
    except Exception:
        return max_cm


def build(md_path, out_path, images_dir):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    lines = Path(md_path).read_text(encoding="utf-8").split("\n")
    i = 0
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()

        if not s:
            i += 1
            continue

        if s == "---":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            i += 1
            continue

        if s.startswith("```"):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            add_code(doc, buf)
            continue

        m = re.match(r"^!\[[^\]]*\]\(([^)]+)\)", s)
        if m:
            img = Path(m.group(1))
            if not img.is_absolute():
                img = (Path(images_dir) / img.name) if images_dir else img
            caption = None
            if i + 1 < len(lines) and lines[i + 1].strip().startswith(">"):
                caption = lines[i + 1].strip().lstrip("> ").strip()
                i += 1
            add_image(doc, img, caption, max_cm=image_size_cm(img))
            i += 1
            continue

        if s.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            rows = []
            header = [c.strip() for c in s.strip("|").split("|")]
            rows.append(header)
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            add_table(doc, rows)
            continue

        hm = re.match(r"^(#{1,4})\s+(.*)$", s)
        if hm:
            style_heading(doc, hm.group(2).strip(), min(len(hm.group(1)), 3))
            i += 1
            continue

        bm = re.match(r"^([-*•])\s+(.*)$", s)
        if bm:
            indent = len(ln) - len(ln.lstrip())
            add_bullet(doc, bm.group(2).strip(), level=1 if indent >= 2 else 0)
            i += 1
            continue

        if s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.4
            add_inline(p, s.lstrip("> ").strip())
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        add_body(doc, s)
        i += 1

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"OK  {out_path}")


if __name__ == "__main__":
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    imgdir = None
    for k, a in enumerate(sys.argv):
        if a == "--images-dir" and k + 1 < len(sys.argv):
            imgdir = sys.argv[k + 1]
    if len(args) < 2:
        print(__doc__)
        sys.exit(1)
    build(args[0], args[1], imgdir)
